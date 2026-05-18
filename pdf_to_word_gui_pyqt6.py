import sys
import os
import site

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLabel, QFileDialog, QProgressBar, QTextEdit,
    QCheckBox, QGroupBox, QLineEdit, QMessageBox
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
import fitz  # PyMuPDF
from pdf2docx import Converter
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'D:\Program Files\Tesseract-OCR\tesseract.exe'
from PIL import Image
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches

import requests
from formula_processor import FormulaProcessor


class PDFConverterThread(QThread):
    progress_update = pyqtSignal(int)
    status_update = pyqtSignal(str)
    finished = pyqtSignal(bool, str)

    def __init__(self, pdf_path, output_path, use_ocr=False, detect_formulas=False):
        super().__init__()
        self.pdf_path = pdf_path
        self.output_path = output_path
        self.use_ocr = use_ocr
        self.detect_formulas = detect_formulas
        self.formula_processor = FormulaProcessor()

    def run(self):
        try:
            if self.detect_formulas:
                self.status_update.emit("Processing formulas...")
                self._convert_with_formula_detection()
            elif self.use_ocr:
                self.status_update.emit("Processing PDF with OCR...")
                self._convert_with_ocr()
            else:
                self.status_update.emit("Converting PDF...")
                self._convert_with_pdf2docx()
            
            self.finished.emit(True, "Conversion completed!")
        except Exception as e:
            self.finished.emit(False, f"Conversion failed: {str(e)}")

    def _convert_with_pdf2docx(self):
        cv = Converter(self.pdf_path)
        cv.convert(self.output_path, start=0, end=None)
        cv.close()

    def _convert_with_ocr(self):
        doc = fitz.open(self.pdf_path)
        word_doc = Document()
        
        total_pages = len(doc)
        
        for page_num, page in enumerate(doc):
            self.progress_update.emit(int((page_num / total_pages) * 100))
            
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            text = pytesseract.image_to_string(img)
            word_doc.add_paragraph(text)
            
            word_doc.add_page_break()
        
        word_doc.save(self.output_path)
        doc.close()

    def _convert_with_formula_detection(self):
        doc = fitz.open(self.pdf_path)
        word_doc = Document()
        
        total_pages = len(doc)
        
        for page_num, page in enumerate(doc):
            self.progress_update.emit(int((page_num / total_pages) * 100))
            
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            img_np = np.array(img)
            has_formulas, formula_regions = self.formula_processor.detect_formulas(img_np)
            
            text = pytesseract.image_to_string(img)
            word_doc.add_paragraph(text)
            
            for region in formula_regions:
                x, y, w, h = region
                formula_img = img.crop((x, y, x + w, y + h))
                temp_path = f"temp_formula.png"
                formula_img.save(temp_path)
                word_doc.add_picture(temp_path, width=Inches(4))
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            
            word_doc.add_page_break()
        
        word_doc.save(self.output_path)
        doc.close()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF to Word Converter")
        self.setGeometry(100, 100, 800, 600)
        
        self.init_ui()

    def init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        layout = QVBoxLayout()
        central_widget.setLayout(layout)
        
        file_group = QGroupBox("File Selection")
        file_layout = QVBoxLayout()
        
        pdf_layout = QHBoxLayout()
        self.pdf_path_label = QLabel("Select PDF file:")
        self.pdf_path_edit = QLineEdit()
        self.pdf_browse_btn = QPushButton("Browse...")
        self.pdf_browse_btn.clicked.connect(self.browse_pdf)
        pdf_layout.addWidget(self.pdf_path_label)
        pdf_layout.addWidget(self.pdf_path_edit)
        pdf_layout.addWidget(self.pdf_browse_btn)
        
        output_layout = QHBoxLayout()
        self.output_path_label = QLabel("Save as:")
        self.output_path_edit = QLineEdit()
        self.output_browse_btn = QPushButton("Browse...")
        self.output_browse_btn.clicked.connect(self.browse_output)
        output_layout.addWidget(self.output_path_label)
        output_layout.addWidget(self.output_path_edit)
        output_layout.addWidget(self.output_browse_btn)
        
        file_layout.addLayout(pdf_layout)
        file_layout.addLayout(output_layout)
        file_group.setLayout(file_layout)
        
        options_group = QGroupBox("Options")
        options_layout = QVBoxLayout()
        
        self.use_ocr_check = QCheckBox("Use OCR (for scanned PDFs)")
        self.detect_formulas_check = QCheckBox("Detect and process formulas")
        options_layout.addWidget(self.use_ocr_check)
        options_layout.addWidget(self.detect_formulas_check)
        options_group.setLayout(options_layout)
        
        self.convert_btn = QPushButton("Convert")
        self.convert_btn.clicked.connect(self.start_conversion)
        
        self.progress_bar = QProgressBar()
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        
        layout.addWidget(file_group)
        layout.addWidget(options_group)
        layout.addWidget(self.convert_btn)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.log_text)

    def browse_pdf(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path_edit.setText(file_path)
            self.output_path_edit.setText(file_path.replace(".pdf", ".docx"))

    def browse_output(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Save As", "", "Word Documents (*.docx)")
        if file_path:
            self.output_path_edit.setText(file_path)

    def start_conversion(self):
        pdf_path = self.pdf_path_edit.text()
        output_path = self.output_path_edit.text()
        
        if not pdf_path or not output_path:
            QMessageBox.warning(self, "Warning", "Please select PDF file and output path!")
            return
        
        if not os.path.exists(pdf_path):
            QMessageBox.warning(self, "Warning", "PDF file not found!")
            return
        
        self.convert_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        self.log_text.clear()
        
        self.thread = PDFConverterThread(
            pdf_path, output_path,
            self.use_ocr_check.isChecked(),
            self.detect_formulas_check.isChecked()
        )
        self.thread.progress_update.connect(self.update_progress)
        self.thread.status_update.connect(self.log_message)
        self.thread.finished.connect(self.conversion_finished)
        self.thread.start()

    def update_progress(self, value):
        self.progress_bar.setValue(value)

    def log_message(self, message):
        self.log_text.append(message)

    def conversion_finished(self, success, message):
        self.convert_btn.setEnabled(True)
        if success:
            QMessageBox.information(self, "Success", message)
        else:
            QMessageBox.warning(self, "Error", message)
        self.log_text.append(message)


def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
