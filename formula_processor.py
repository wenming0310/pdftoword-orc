import fitz  # PyMuPDF
import numpy as np
from PIL import Image
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r'D:\Program Files\Tesseract-OCR\tesseract.exe'
from docx import Document
from docx.shared import Inches
import re
import cv2


class FormulaProcessor:
    def __init__(self):
        self.math_symbols = [
            '=', '+', '-', '×', '÷', '∫', '∑', '∏', '√', '∞',
            '≠', '≤', '≥', '∈', '∪', '∩', '→', '←', '↔', '⇒',
            '∀', '∂', 'Δ', '∇', '⊂', '⊃', '⊆', '⊇', 'α', 'β',
            'γ', 'δ', 'ε', 'θ', 'λ', 'μ', 'π', 'σ', 'φ', 'ω',
            '()', '[]', '{}', 'frac', 'sqrt', 'sum', 'int'
        ]
    
    def detect_formulas(self, img_np):
        """
        Advanced formula detection using multiple methods
        Returns (has_formulas: bool, formula_regions: list of (x, y, w, h))
        """
        formula_regions = []
        
        # Method 1: Use Tesseract to detect text blocks and analyze
        text_regions = self._detect_text_blocks_ocr(img_np)
        
        for region in text_regions:
            x, y, w, h, text = region
            if self._is_formula_text(text):
                # Expand the region slightly for better capture
                expand_x = max(0, x - 10)
                expand_y = max(0, y - 5)
                expand_w = w + 20
                expand_h = h + 10
                formula_regions.append((expand_x, expand_y, expand_w, expand_h))
        
        # Method 2: Detect isolated mathematical expressions using image processing
        image_regions = self._detect_formula_by_image(img_np)
        formula_regions.extend(image_regions)
        
        # Remove duplicates and merge overlapping regions
        formula_regions = self._merge_regions(formula_regions)
        
        has_formulas = len(formula_regions) > 0
        return has_formulas, formula_regions
    
    def _detect_text_blocks_ocr(self, img_np):
        """Use Tesseract OCR to detect text blocks"""
        try:
            img_pil = Image.fromarray(img_np)
            data = pytesseract.image_to_data(img_pil, output_type=pytesseract.Output.DICT)
            
            blocks = []
            n_boxes = len(data['text'])
            
            i = 0
            while i < n_boxes:
                if int(data['conf'][i]) > 30:  # Confidence threshold
                    text = data['text'][i].strip()
                    if text:
                        x = data['left'][i]
                        y = data['top'][i]
                        w = data['width'][i]
                        h = data['height'][i]
                        
                        # Group nearby text blocks
                        block_text = text
                        j = i + 1
                        while j < n_boxes and int(data['conf'][j]) > 30:
                            next_text = data['text'][j].strip()
                            if next_text:
                                # Check if text is on same line
                                if abs(data['top'][j] - y) < h:
                                    block_text += ' ' + next_text
                                    w = max(w, data['left'][j] + data['width'][j] - x)
                                    j += 1
                                else:
                                    break
                            else:
                                j += 1
                        
                        if block_text.strip():
                            blocks.append((x, y, w, h, block_text))
                        i = j
                    else:
                        i += 1
                else:
                    i += 1
            
            return blocks
        except:
            return []
    
    def _is_formula_text(self, text):
        """Check if text is likely a formula"""
        if not text or len(text.strip()) < 2:
            return False
        
        # Check for math symbols
        math_count = sum(1 for symbol in self.math_symbols if symbol in text)
        
        # Check for numbers and variables pattern
        has_numbers = any(char.isdigit() for char in text)
        has_operators = any(op in text for op in ['+', '-', '×', '÷', '=', '/', '*'])
        
        # Check for fraction pattern (a/b)
        has_fraction = bool(re.search(r'\d+/\d+', text))
        
        # Check for superscript/subscript indicators
        has_exponent = bool(re.search(r'\d+\^|\^\d+|\d+\*\*', text))
        
        # Check for Greek letters
        has_greek = bool(re.search(r'[αβγδεζηθικλμνξοπρστυφχψω]', text, re.IGNORECASE))
        
        # Check for special LaTeX-like patterns
        has_latex = bool(re.search(r'\\frac|\\sqrt|\\sum|\\int|\\alpha|\\beta', text, re.IGNORECASE))
        
        # Check if line is mostly math (high ratio of math characters)
        math_chars = sum(1 for c in text if c.isdigit() or c in '+-*/=<>≥≤≠()[]{}')
        math_ratio = math_chars / len(text) if len(text) > 0 else 0
        
        # Decision rules
        if math_count >= 2:
            return True
        if has_fraction or has_exponent:
            return True
        if has_greek and (has_numbers or has_operators):
            return True
        if has_latex:
            return True
        if math_ratio > 0.4 and has_numbers and has_operators:
            return True
        if math_count >= 1 and has_numbers and len(text) < 30:
            return True
        
        return False
    
    def _detect_formula_by_image(self, img_np):
        """Detect formulas using image processing techniques"""
        formula_regions = []
        
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        
        # Apply adaptive thresholding
        binary = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY_INV, 11, 2
        )
        
        # Detect horizontal lines (formulas often centered)
        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
        detect_horizontal = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel)
        
        # Find contours
        contours, _ = cv2.findContours(detect_horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        img_height, img_width = img_np.shape[:2]
        
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            
            # Formula lines are usually centered and wide
            center_x = x + w / 2
            is_centered = abs(center_x - img_width / 2) < img_width * 0.3
            
            # Wide enough to be a formula
            is_wide = w > img_width * 0.3
            
            # Not too tall
            is_not_too_tall = h < img_height * 0.1
            
            if is_centered and is_wide and is_not_too_tall:
                formula_regions.append((x, y - 5, w, h + 10))
        
        # Also detect isolated blocks using vertical projection
        vert_projection = np.sum(binary, axis=0)
        threshold = np.mean(vert_projection) * 0.1
        
        in_formula = False
        formula_start = 0
        
        for i, val in enumerate(vert_projection):
            if val < threshold and not in_formula:
                in_formula = True
                formula_start = i
            elif val >= threshold and in_formula:
                in_formula = False
                region_width = i - formula_start
                if region_width > 100:
                    formula_regions.append((formula_start, 0, region_width, img_height))
        
        return formula_regions
    
    def _merge_regions(self, regions):
        """Merge overlapping or nearby regions"""
        if not regions:
            return []
        
        # Sort by x coordinate
        regions = sorted(regions, key=lambda r: r[0])
        
        merged = []
        current = list(regions[0])
        
        for region in regions[1:]:
            # Check if overlapping or nearby (within 20 pixels)
            if region[0] <= current[0] + current[2] + 20:
                # Merge by expanding
                current[0] = min(current[0], region[0])
                current[1] = min(current[1], region[1])
                current[2] = max(current[0] + current[2], region[0] + region[2]) - current[0]
                current[3] = max(current[1] + current[3], region[1] + region[3]) - current[1]
            else:
                merged.append(tuple(current))
                current = list(region)
        
        merged.append(tuple(current))
        
        # Filter out very small regions
        filtered = [r for r in merged if r[2] > 50 and r[3] > 20]
        
        return filtered

    def is_likely_formula(self, text):
        """判断文本是否可能是公式"""
        math_symbols = 0
        for keyword in self.math_symbols:
            math_symbols += text.count(keyword)
        
        has_numbers = any(char.isdigit() for char in text)
        has_letters = any(char.isalpha() for char in text)
        
        return math_symbols > 0 or (has_numbers and has_letters and len(text) > 5)

    def extract_formulas_from_pdf(self, pdf_path):
        """从PDF中提取可能包含公式的区域"""
        formulas = []
        pdf = fitz.open(pdf_path)
        
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            blocks = page.get_text("blocks")
            
            for block in blocks:
                x0, y0, x1, y1, text, block_no, block_type = block
                
                if self.is_likely_formula(text):
                    formulas.append({
                        'page': page_num,
                        'bbox': (x0, y0, x1, y1),
                        'text': text,
                        'position': (x0, y0)
                    })
        
        pdf.close()
        return formulas

    def process_pdf_with_formulas(self, pdf_path, output_path, use_ocr=False):
        """处理包含公式的PDF"""
        doc = Document()
        pdf = fitz.open(pdf_path)
        
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            if use_ocr:
                text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                if text.strip():
                    doc.add_paragraph(text)
            else:
                text = page.get_text()
                if text.strip():
                    doc.add_paragraph(text)
            
            if page_num < len(pdf) - 1:
                doc.add_page_break()
        
        pdf.close()
        doc.save(output_path)
        return True
