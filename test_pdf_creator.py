
from docx import Document
from docx.shared import Inches
import os

# 创建一个测试Word文档，然后我们可以尝试转换它
def create_test_word():
    doc = Document()
    doc.add_heading('PDF转Word测试文档', 0)
    doc.add_paragraph('这是一个用于测试PDF转Word功能的文档。')
    doc.add_heading('数学公式测试', level=1)
    doc.add_paragraph('常见数学公式示例：')
    doc.add_paragraph('1 + 1 = 2')
    doc.add_paragraph('E = mc²')
    doc.add_paragraph('∑_{i=1}^{n} i = n(n+1)/2')
    doc.add_paragraph('∫_{0}^{∞} e^{-x²} dx = √π / 2')
    doc.add_page_break()
    doc.add_heading('普通文本测试', level=1)
    doc.add_paragraph('这是一段普通的中文文本，用于测试OCR功能是否能正常识别中文字符。')
    doc.add_paragraph('这是一段普通的英文文本，用于测试OCR功能是否能正常识别英文字符。')
    test_doc_path = '/workspace/test_document.docx'
    doc.save(test_doc_path)
    print(f'测试文档已创建: {test_doc_path}')
    return test_doc_path

if __name__ == '__main__':
    create_test_word()

