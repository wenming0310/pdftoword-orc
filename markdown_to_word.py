from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re


class MarkdownToWord:
    def __init__(self, markdown_text, image_dir='temp_images'):
        self.markdown_text = markdown_text
        self.image_dir = image_dir
        self.doc = Document()
        
        # Set default style
        self._setup_document_style()
    
    def _setup_document_style(self):
        """Setup document default styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def convert(self, output_path):
        """Convert Markdown to Word document"""
        lines = self.markdown_text.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            
            if line.startswith('# '):
                # Heading 1
                self._add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # Heading 2
                self._add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # Heading 3
                self._add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                # Heading 4
                self._add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                # Heading 5
                self._add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                # Heading 6
                self._add_heading(line[7:], level=6)
            elif line.startswith('!['):
                # Image
                self._add_image(line)
            elif line.startswith('<!--'):
                # Comment (skip)
                pass
            elif line.startswith('---') or line.startswith('***'):
                # Horizontal rule
                self._add_horizontal_rule()
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                i = self._add_list(lines, i, 'bullet')
                continue
            elif re.match(r'^\d+\. ', line):
                # Numbered list
                i = self._add_list(lines, i, 'number')
                continue
            elif line.strip() == '':
                # Empty line (add small space)
                self.doc.add_paragraph()
            elif line.startswith('> '):
                # Block quote
                self._add_quote(line[2:])
            elif line.startswith('```'):
                # Code block
                i = self._add_code_block(lines, i)
                continue
            else:
                # Regular text
                if line.strip():
                    self._add_paragraph(line)
            
            i += 1
        
        self.doc.save(output_path)
    
    def _add_heading(self, text, level):
        """Add heading to document"""
        self.doc.add_heading(text, level=level)
    
    def _add_paragraph(self, text):
        """Add paragraph to document"""
        paragraph = self.doc.add_paragraph()
        
        # Handle inline formatting (bold, italic, etc.)
        self._add_formatted_text(paragraph, text)
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with inline formatting"""
        # Simple implementation for inline formatting
        # This handles **bold**, *italic*, and `code`
        
        # Split by formatting markers
        i = 0
        n = len(text)
        
        while i < n:
            if text[i:i+2] == '**' and i+2 < n:
                # Bold start
                end = text.find('**', i+2)
                if end != -1:
                    run = paragraph.add_run(text[i+2:end])
                    run.bold = True
                    i = end + 2
                else:
                    paragraph.add_run(text[i])
                    i += 1
            elif text[i] == '*' and i+1 < n and text[i+1] != '*':
                # Italic start
                end = text.find('*', i+1)
                if end != -1:
                    run = paragraph.add_run(text[i+1:end])
                    run.italic = True
                    i = end + 1
                else:
                    paragraph.add_run(text[i])
                    i += 1
            elif text[i] == '`' and i+1 < n:
                # Code start
                end = text.find('`', i+1)
                if end != -1:
                    run = paragraph.add_run(text[i+1:end])
                    run.font.name = 'Courier New'
                    i = end + 1
                else:
                    paragraph.add_run(text[i])
                    i += 1
            else:
                paragraph.add_run(text[i])
                i += 1
    
    def _add_image(self, line):
        """Add image to document"""
        # Parse ![]() format
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text, img_path = match.groups()
            
            # Try to load the image
            if os.path.exists(img_path):
                try:
                    # Add image with appropriate size
                    from PIL import Image as PILImage
                    img = PILImage.open(img_path)
                    width, height = img.size
                    
                    # Calculate appropriate size (max 16 cm width)
                    max_width_cm = 16
                    aspect_ratio = height / width
                    
                    # Convert pixels to inches (approx)
                    dpi = 96
                    width_in = width / dpi
                    height_in = height / dpi
                    
                    if width_in > max_width_cm / 2.54:
                        scale = (max_width_cm / 2.54) / width_in
                        width_in *= scale
                        height_in *= scale
                    
                    # Add image
                    paragraph = self.doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(img_path, width=Inches(width_in))
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
                    # Fallback to alt text
                    self.doc.add_paragraph(f"[Image: {alt_text}]")
            else:
                # Image not found, add alt text
                self.doc.add_paragraph(f"[Image: {alt_text}]")
    
    def _add_horizontal_rule(self):
        """Add horizontal rule"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run('─' * 40)
    
    def _add_quote(self, text):
        """Add block quote"""
        paragraph = self.doc.add_paragraph(text)
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.right_indent = Inches(0.5)
    
    def _add_code_block(self, lines, start_idx):
        """Add code block"""
        end_idx = start_idx + 1
        while end_idx < len(lines) and not lines[end_idx].startswith('```'):
            end_idx += 1
        
        # Add code lines
        code_lines = lines[start_idx + 1:end_idx]
        for line in code_lines:
            paragraph = self.doc.add_paragraph(line)
            paragraph.paragraph_format.left_indent = Inches(0.5)
            run = paragraph.runs[0]
            run.font.name = 'Courier New'
        
        return end_idx
    
    def _add_list(self, lines, start_idx, list_type):
        """Add bullet or numbered list"""
        i = start_idx
        
        if list_type == 'bullet':
            while i < len(lines):
                line = lines[i].rstrip()
                if line.startswith('- ') or line.startswith('* '):
                    self.doc.add_paragraph(line[2:], style='List Bullet')
                    i += 1
                else:
                    break
        elif list_type == 'number':
            while i < len(lines):
                if re.match(r'^\d+\. ', lines[i]):
                    self.doc.add_paragraph(re.sub(r'^\d+\. ', '', lines[i]), style='List Number')
                    i += 1
                else:
                    break
        
        return i - 1
